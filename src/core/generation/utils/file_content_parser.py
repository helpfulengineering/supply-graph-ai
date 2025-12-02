"""
File Content Parser utility for file categorization.

This module provides utilities for parsing file content at different
analysis depths for LLM-based file categorization.

Supported file formats:
- Text files: .md, .txt, .rst (direct content extraction)
- PDF files: .pdf (text extraction using pypdf)
- DOCX files: .docx (text extraction using python-docx)
- Other binary files: Not supported (returns None)

Note: Old .doc format is not supported. Only .docx format is supported.
"""

import logging
import re
from pathlib import Path
from typing import Optional

from ..models import AnalysisDepth, FileInfo

logger = logging.getLogger(__name__)


class FileContentParser:
    """
    Utility for parsing file content based on analysis depth.

    This class provides methods to extract file content at different
    depth levels (Shallow/Medium/Deep) for LLM-based categorization.
    It handles both text files and binary files (with text extraction).
    """

    # Depth limits in characters
    SHALLOW_LIMIT = 500
    MEDIUM_LIMIT = 2000

    def parse_content(self, file_info: FileInfo, depth: AnalysisDepth) -> Optional[str]:
        """
        Parse file content based on analysis depth.

        Args:
            file_info: File to parse
            depth: Analysis depth level (Shallow/Medium/Deep)

        Returns:
            Parsed content string or None if not parseable

        Behavior:
            - Text files (.md, .txt, .rst): Extract content based on depth
            - Binary files (PDF, DOCX): Attempt text extraction, return None if fails
            - Logs skipped files for visibility
        """
        # Handle None content
        if file_info.content is None:
            # Try binary extraction for binary file types
            if self._is_binary_file(file_info):
                extracted_text = self.extract_binary_text(file_info)
                if extracted_text is None:
                    return None
                # Apply depth limits to extracted text
                if depth == AnalysisDepth.SHALLOW:
                    return extracted_text[: self.SHALLOW_LIMIT]
                elif depth == AnalysisDepth.MEDIUM:
                    return extracted_text[: self.MEDIUM_LIMIT]
                elif depth == AnalysisDepth.DEEP:
                    return extracted_text
                else:
                    # Default to shallow
                    return extracted_text[: self.SHALLOW_LIMIT]
            return None

        # Handle empty content
        if not file_info.content:
            return ""

        # Apply depth limits
        if depth == AnalysisDepth.SHALLOW:
            return file_info.content[: self.SHALLOW_LIMIT]
        elif depth == AnalysisDepth.MEDIUM:
            return file_info.content[: self.MEDIUM_LIMIT]
        elif depth == AnalysisDepth.DEEP:
            return file_info.content
        else:
            # Default to shallow
            logger.warning(f"Unknown depth {depth}, defaulting to shallow")
            return file_info.content[: self.SHALLOW_LIMIT]

    def extract_binary_text(self, file_info: FileInfo) -> Optional[str]:
        """
        Attempt to extract text from binary files.

        Args:
            file_info: Binary file to extract text from

        Returns:
            Extracted text or None if extraction fails

        Supported formats:
            - PDF: Extract text content using pypdf
            - DOCX: Extract text content using python-docx
            - Other binary: Returns None, logs attempt
        """
        file_path = Path(file_info.path)
        file_ext = file_path.suffix.lower()

        # Resolve file path (handle both absolute and relative)
        try:
            if not file_path.is_absolute():
                # If relative, try to resolve from current working directory
                resolved_path = file_path.resolve()
            else:
                resolved_path = file_path

            # Check if file exists
            if not resolved_path.exists():
                logger.warning(f"File not found for text extraction: {resolved_path}")
                return None

            # Check file size (skip very large files to avoid memory issues)
            file_size = resolved_path.stat().st_size
            max_file_size = 50 * 1024 * 1024  # 50 MB limit
            if file_size > max_file_size:
                logger.warning(
                    f"File too large for text extraction: {resolved_path} ({file_size} bytes)"
                )
                return None

        except Exception as e:
            logger.warning(
                f"Error resolving file path {file_info.path}: {e}", exc_info=True
            )
            return None

        # PDF extraction
        if file_ext == ".pdf":
            return self._extract_pdf_text(resolved_path)

        # DOCX extraction
        if file_ext == ".docx":
            return self._extract_docx_text(resolved_path)

        # DOC (old format) - not supported yet
        if file_ext == ".doc":
            logger.debug(
                f"Old .doc format not supported for text extraction: {file_info.path}. "
                "Only .docx format is supported."
            )
            return None

        # Other binary formats - not supported yet
        logger.debug(
            f"Binary file type {file_ext} not supported for text extraction: {file_info.path}"
        )
        return None

    def _is_binary_file(self, file_info: FileInfo) -> bool:
        """
        Check if file is a binary file type.

        Args:
            file_info: File to check

        Returns:
            True if file is binary, False otherwise
        """
        file_path = Path(file_info.path)
        file_ext = file_path.suffix.lower()

        # Binary file extensions
        binary_extensions = {
            ".pdf",
            ".docx",
            ".doc",
            ".xlsx",
            ".xls",
            ".pptx",
            ".ppt",
            ".zip",
            ".tar",
            ".gz",
            ".rar",
            ".7z",
            ".jpg",
            ".jpeg",
            ".png",
            ".gif",
            ".bmp",
            ".svg",
            ".mp3",
            ".mp4",
            ".avi",
            ".mov",
            ".exe",
            ".dll",
            ".so",
            ".dylib",
            ".bin",
            ".dat",
        }

        # Check file type
        binary_file_types = {"pdf", "image", "binary", "archive", "video", "audio"}

        return file_ext in binary_extensions or file_info.file_type in binary_file_types

    def _extract_pdf_text(self, file_path: Path) -> Optional[str]:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            import pypdf

            text_parts = []

            with open(file_path, "rb") as pdf_file:
                try:
                    pdf_reader = pypdf.PdfReader(pdf_file)

                    # Check if PDF is encrypted
                    if pdf_reader.is_encrypted:
                        logger.warning(
                            f"PDF is encrypted, cannot extract text: {file_path}"
                        )
                        return None

                    # Extract text from each page
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                        except Exception as e:
                            logger.warning(
                                f"Error extracting text from PDF page {page_num + 1} "
                                f"in {file_path}: {e}"
                            )
                            # Continue with other pages
                            continue

                    if not text_parts:
                        logger.debug(f"No text found in PDF: {file_path}")
                        return None

                    # Combine all pages with double newline separator
                    extracted_text = "\n\n".join(text_parts)

                    # Clean up excessive whitespace
                    extracted_text = self._clean_extracted_text(extracted_text)

                    logger.debug(
                        f"Successfully extracted {len(extracted_text)} characters "
                        f"from PDF: {file_path}"
                    )

                    return extracted_text

                except pypdf.errors.PdfReadError as e:
                    logger.warning(
                        f"Error reading PDF file {file_path}: {e}", exc_info=True
                    )
                    return None
                except Exception as e:
                    logger.warning(
                        f"Unexpected error extracting text from PDF {file_path}: {e}",
                        exc_info=True,
                    )
                    return None

        except ImportError:
            logger.warning(
                "pypdf library not installed. Install with: pip install pypdf"
            )
            return None
        except Exception as e:
            logger.warning(
                f"Failed to extract text from PDF {file_path}: {e}", exc_info=True
            )
            return None

    def _extract_docx_text(self, file_path: Path) -> Optional[str]:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            from docx import Document

            text_parts = []

            try:
                doc = Document(file_path)

                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    para_text = paragraph.text.strip()
                    if para_text:
                        text_parts.append(para_text)

                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_texts.append(cell_text)
                        if row_texts:
                            text_parts.append(" | ".join(row_texts))

                if not text_parts:
                    logger.debug(f"No text found in DOCX: {file_path}")
                    return None

                # Combine all text parts with double newline separator
                extracted_text = "\n\n".join(text_parts)

                # Clean up excessive whitespace
                extracted_text = self._clean_extracted_text(extracted_text)

                logger.debug(
                    f"Successfully extracted {len(extracted_text)} characters "
                    f"from DOCX: {file_path}"
                )

                return extracted_text

            except Exception as e:
                logger.warning(
                    f"Error reading DOCX file {file_path}: {e}", exc_info=True
                )
                return None

        except ImportError:
            logger.warning(
                "python-docx library not installed. Install with: pip install python-docx"
            )
            return None
        except Exception as e:
            logger.warning(
                f"Failed to extract text from DOCX {file_path}: {e}", exc_info=True
            )
            return None

    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean up extracted text by removing excessive whitespace and corrupted patterns.

        This method filters out common PDF extraction artifacts:
        - Control characters and non-printable characters
        - Corrupted text patterns (random alphanumeric sequences)
        - Excessive whitespace

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        # Remove control characters and non-printable characters (except newlines and tabs)
        # This helps filter out corrupted text from PDF extraction
        text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]", "", text)

        # Filter out lines that look like corrupted text patterns
        # Pattern: random alphanumeric sequences with underscores (e.g., "h28qPJWAI_3NuLXk1RQ_")
        lines = text.split("\n")
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                cleaned_lines.append("")
                continue

            # Check for random-looking patterns that suggest PDF corruption
            # Skip lines that are mostly random alphanumeric with underscores
            if re.match(r"^[a-zA-Z0-9_]+$", line) and "_" in line and len(line) > 8:
                # Check if it looks like a random sequence
                unique_chars = len(set(line.lower()))
                if unique_chars > len(line) * 0.6:  # High character diversity
                    parts = line.split("_")
                    if len(parts) >= 3:
                        part_lengths = [len(p) for p in parts]
                        avg_length = sum(part_lengths) / len(part_lengths)
                        variance = sum(
                            (l - avg_length) ** 2 for l in part_lengths
                        ) / len(part_lengths)
                        if variance > 10:  # High variance suggests random pattern
                            # Skip this line - it's likely corrupted
                            continue

            # Check for excessive case changes (suggests corruption)
            case_changes = 0
            alpha_chars = 0
            for i in range(len(line) - 1):
                if line[i].isalpha() and line[i + 1].isalpha():
                    alpha_chars += 1
                    if line[i].islower() != line[i + 1].islower():
                        case_changes += 1

            if alpha_chars > 0 and case_changes > alpha_chars * 0.4:
                words = line.split()
                if len(words) == 1 and case_changes > 3:
                    # Single word with many case changes - likely corrupted
                    continue
                if len(words) > 1 and case_changes > 5:
                    # Multiple words but still high case change ratio - likely corrupted
                    continue

            cleaned_lines.append(line)

        text = "\n".join(cleaned_lines)

        # Replace multiple spaces with single space
        text = re.sub(r" +", " ", text)

        # Replace multiple newlines (3+) with double newline
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()
