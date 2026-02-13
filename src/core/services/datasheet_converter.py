"""
MSF Datasheet Converter Service for Open Hardware Manager (OHM).

Provides bi-directional conversion between OKH manifests and MSF
(Maker Space Foundation) 3D-printed product technical specification
datasheet format (.docx).

The OKH canonical data model is always the source of truth. The converter
maps OKH fields to/from the structured tables in the MSF datasheet
template.

MSF Datasheet Structure (6 tables):
  Table 0 - Header (logo + title)
  Table 1 - Identity (internal ref, name, product stage)
  Table 2 - FORM (picture, category, dimensions, license, readiness)
  Table 3 - FIT (compatibility, manufacturing instructions, materials)
  Table 4 - FUNCTION (description, cleaning, packaging, standards)
  Table 5 - ATTACHMENTS (drawings, designed/approved/tested by)

Usage:
    converter = DatasheetConverter()
    # OKH → docx
    converter.okh_to_datasheet(manifest, output_path)
    # docx → OKH
    manifest = converter.datasheet_to_okh(docx_path)
"""

import os
import re
from datetime import date
from typing import Any, Dict, List, Optional

from ..models.okh import (
    DocumentRef,
    DocumentationType,
    License,
    MaterialSpec,
    OKHManifest,
    Person,
    Standard,
)
from ..utils.logging import get_logger

logger = get_logger(__name__)

# Default template location relative to the project root
_DEFAULT_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__),
    "..",
    "..",
    "..",
    "notes",
    "msf",
    "datasheet-template.docx",
)


class DatasheetConversionError(Exception):
    """Raised when datasheet conversion fails."""

    pass


class DatasheetConverter:
    """Bi-directional converter between OKH manifests and MSF datasheets.

    The converter uses the MSF 3D-printed product technical specification
    datasheet template as the target/source docx format.  The internal,
    canonical OKHManifest dataclass is always the authoritative
    representation.
    """

    def __init__(self, template_path: Optional[str] = None):
        """Initialise the converter.

        Args:
            template_path: Path to the MSF .docx template.  When *None*
                the built-in template shipped with OHM is used.
        """
        self.template_path = template_path or os.path.normpath(_DEFAULT_TEMPLATE_PATH)
        if not os.path.exists(self.template_path):
            raise DatasheetConversionError(
                f"Datasheet template not found at: {self.template_path}"
            )

    # ------------------------------------------------------------------
    # OKH → Datasheet (.docx)
    # ------------------------------------------------------------------

    def okh_to_datasheet(
        self,
        manifest: OKHManifest,
        output_path: str,
    ) -> str:
        """Convert an OKHManifest to a populated MSF datasheet.

        Args:
            manifest: The canonical OKH manifest to convert.
            output_path: Filesystem path for the generated .docx file.

        Returns:
            The absolute path of the written file.

        Raises:
            DatasheetConversionError: If the conversion or write fails.
        """
        try:
            from docx import Document
        except ImportError:
            raise DatasheetConversionError(
                "python-docx is required for datasheet conversion. "
                "Install it with: pip install python-docx"
            )

        try:
            doc = Document(self.template_path)
        except Exception as exc:
            raise DatasheetConversionError(f"Failed to open template: {exc}") from exc

        tables = doc.tables
        if len(tables) < 6:
            raise DatasheetConversionError(
                f"Template has {len(tables)} tables, expected at least 6."
            )

        # --- Table 1: Identity -------------------------------------------
        self._set_cell(
            tables[1],
            0,
            0,
            f"Internal Ref.\n{manifest.metadata.get('internal_ref', str(manifest.id)[:8])}",
        )
        self._set_cell(tables[1], 0, 1, f"Name:\n{manifest.title}")
        self._set_cell(
            tables[1], 0, 2, f"Product stage:\n{manifest.development_stage or ''}"
        )

        # --- Table 2: FORM -----------------------------------------------
        self._set_cell(tables[2], 0, 0, f"Product picture:\n{manifest.image or ''}")

        category = ""
        subcategory = ""
        if manifest.keywords:
            category = manifest.keywords[0] if len(manifest.keywords) > 0 else ""
            subcategory = manifest.keywords[1] if len(manifest.keywords) > 1 else ""

        meta = manifest.metadata or {}
        form_detail = (
            f"Category: {category}\n"
            f"Subcategory: {subcategory}\n"
            f"Critical item: {meta.get('critical_item', '')}\n"
            f"Dangerous goods: {meta.get('dangerous_goods', '')}\n"
            f"Short description: {manifest.description or ''}"
        )
        self._set_cell(tables[2], 0, 1, form_detail)

        # Repository link
        self._set_cell(tables[2], 1, 1, f"Repository link: {manifest.repo or ''}")

        # Dimensions & use type
        dims = ""
        if (
            manifest.manufacturing_specs
            and manifest.manufacturing_specs.outer_dimensions
        ):
            d = manifest.manufacturing_specs.outer_dimensions
            dims = f"{d.get('length', '')}x{d.get('width', '')}x{d.get('height', '')}"

        dim_cell = (
            f"Overall dimensions:\nLxWxH: {dims}\n"
            f"Single/Multiple use: {meta.get('use_type', '')}\n"
            f"Permanent/Temporary solution: {meta.get('solution_type', '')}"
        )
        self._set_cell(tables[2], 2, 0, dim_cell)

        # License
        license_str = self._license_to_string(manifest.license)
        self._set_cell(tables[2], 2, 1, f"License: {license_str}")

        # Readiness levels
        readiness = (
            f"Field readiness: {meta.get('field_readiness', '')}\n"
            f"Maker readiness: {meta.get('maker_readiness', '')}\n"
            f"User readiness: {meta.get('user_readiness', '')}\n"
            f"Technology readiness: {manifest.technology_readiness_level or ''}\n"
            f"Risk level: {meta.get('risk_level', '')}"
        )
        self._set_cell(tables[2], 3, 1, readiness)

        # Justification
        justification = manifest.intended_use or meta.get("justification", "")
        self._set_cell(
            tables[2], 4, 0, f"Justification of using 3D printed item:\n{justification}"
        )
        self._set_cell(
            tables[2], 4, 1, f"Justification of using 3D printed item:\n{justification}"
        )

        # Why 3D printed / approval
        why_3d = meta.get("justification_3d_print", "")
        approval = meta.get("approval_required_by", "")
        self._set_cell(
            tables[2],
            5,
            0,
            f"Why should this item be 3D printed:\n{why_3d}\nApproval required by: {approval}",
        )
        self._set_cell(
            tables[2],
            5,
            1,
            f"Why should this item be 3D printed:\n{why_3d}\nApproval required by: {approval}",
        )

        # --- Table 3: FIT ------------------------------------------------
        compat = meta.get("primary_compatibility", "")
        accessories = meta.get("compatible_accessories", "")
        self._set_cell(
            tables[3],
            0,
            0,
            f"Primary compatibility: {compat}\n\nCompatible accessories: {accessories}",
        )

        # Manufacturing instructions
        mfg_instr = (
            "; ".join(d.title for d in manifest.making_instructions)
            if manifest.making_instructions
            else ""
        )
        self._set_cell(tables[3], 1, 0, f"Manufacturing instructions: {mfg_instr}")

        # Materials and manufacturing detail
        material_str = (
            ", ".join(m.name for m in manifest.materials) if manifest.materials else ""
        )
        printer = ", ".join(manifest.tool_list) if manifest.tool_list else ""
        slicer = meta.get("slicer_settings", "")
        post_proc = meta.get("post_processing", "")
        assembly = meta.get("assembly_instructions", "")
        qc = meta.get("qc_procedures", "")
        visual_insp = meta.get("visual_inspection", "")
        dim_valid = meta.get("dimensional_validation", "")
        tol_insp = meta.get("tolerance_inspection", "")
        safety_valid = meta.get("safety_validation", "")
        check_freq = meta.get("check_frequency", "")

        fit_detail = (
            f"Material and color: {material_str}\n"
            f"List of other materials: {meta.get('other_materials', '')}\n"
            f"3D Printer: {printer}\n"
            f"Slicer settings: {slicer}\n"
            f"Post processing instructions: {post_proc}\n"
            f"Assembly instructions: {assembly}\n"
            f"QC procedures:\n"
            f"Visual inspection: {visual_insp}\n"
            f"Dimensional validation: {dim_valid}\n"
            f"Tolerance inspection: {tol_insp}\n"
            f"Safety validation: {safety_valid}\n"
            f"Regular product check frequency and responsibility: {check_freq}"
        )
        self._set_cell(tables[3], 2, 0, fit_detail)

        # --- Table 4: FUNCTION -------------------------------------------
        self._set_cell(
            tables[4],
            0,
            0,
            f"Detailed description of the component/product/workflow and its use:\n{manifest.function or ''}",
        )

        # Cleaning
        cleaning = self._get_doc_ref_text(manifest.operating_instructions, "cleaning")
        if not cleaning:
            cleaning = meta.get("cleaning_procedures", "")
        self._set_cell(
            tables[4], 2, 0, f"Cleaning and sanitization procedures: {cleaning}"
        )

        # Packaging
        packaging = self._get_doc_ref_text(manifest.operating_instructions, "packaging")
        if not packaging:
            packaging = meta.get("packaging_instructions", "")
        self._set_cell(
            tables[4], 4, 0, f"Packaging and storing instructions: {packaging}"
        )

        # Standards & safety
        standards_str = ""
        if manifest.standards_used:
            standards_str = "; ".join(
                s.standard_title + (f" ({s.reference})" if s.reference else "")
                for s in manifest.standards_used
            )
        safety = manifest.health_safety_notice or ""
        self._set_cell(
            tables[4],
            6,
            0,
            f"Related links, standards, safety considerations: {standards_str} {safety}".strip(),
        )

        spaulding = meta.get("spaulding_classification", "")
        self._set_cell(tables[4], 7, 0, f"Spaulding Classification (IPC): {spaulding}")

        # --- Table 5: ATTACHMENTS ----------------------------------------
        design_str = (
            "; ".join(d.title for d in manifest.design_files)
            if manifest.design_files
            else ""
        )
        self._set_cell(
            tables[5],
            0,
            0,
            f"Technical drawings\nPhotos: {manifest.image or ''}\nOther: {design_str}",
        )

        # Designed by
        designer = self._licensor_to_string(manifest.licensor)
        self._set_cell(tables[5], 1, 0, f"Designed by\n{designer}")
        date_str = manifest.version_date.isoformat() if manifest.version_date else ""
        self._set_cell(tables[5], 1, 2, f"Date: {date_str}")

        # Approved / Tested
        self._set_cell(
            tables[5], 2, 0, f"Product approved by\n{meta.get('approved_by', '')}"
        )
        self._set_cell(tables[5], 2, 2, f"Date: {meta.get('approved_date', '')}")
        self._set_cell(
            tables[5], 3, 0, f"Product tested by\n{meta.get('tested_by', '')}"
        )
        self._set_cell(tables[5], 3, 2, f"Date: {meta.get('tested_date', '')}")

        # Write output
        try:
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            doc.save(output_path)
            logger.info(f"Datasheet written to {output_path}")
            return os.path.abspath(output_path)
        except Exception as exc:
            raise DatasheetConversionError(f"Failed to save datasheet: {exc}") from exc

    # ------------------------------------------------------------------
    # Datasheet (.docx) → OKH
    # ------------------------------------------------------------------

    def datasheet_to_okh(self, docx_path: str) -> OKHManifest:
        """Parse a filled MSF datasheet and return a canonical OKHManifest.

        Args:
            docx_path: Path to the populated .docx datasheet.

        Returns:
            An OKHManifest populated from the datasheet fields.

        Raises:
            DatasheetConversionError: If the file cannot be parsed.
        """
        try:
            from docx import Document
        except ImportError:
            raise DatasheetConversionError(
                "python-docx is required for datasheet conversion. "
                "Install it with: pip install python-docx"
            )

        if not os.path.exists(docx_path):
            raise DatasheetConversionError(f"File not found: {docx_path}")

        try:
            doc = Document(docx_path)
        except Exception as exc:
            raise DatasheetConversionError(f"Failed to open docx file: {exc}") from exc

        tables = doc.tables
        if len(tables) < 6:
            raise DatasheetConversionError(
                f"Document has {len(tables)} tables, expected at least 6."
            )

        metadata: Dict[str, Any] = {}

        # --- Table 1: Identity -------------------------------------------
        ref_text = self._get_cell(tables[1], 0, 0)
        internal_ref = self._extract_value(ref_text, "Internal Ref.")
        if internal_ref:
            metadata["internal_ref"] = internal_ref

        title = self._extract_value(self._get_cell(tables[1], 0, 1), "Name:")
        dev_stage = self._extract_value(
            self._get_cell(tables[1], 0, 2), "Product stage:"
        )

        # --- Table 2: FORM -----------------------------------------------
        image_text = self._get_cell(tables[2], 0, 0)
        image = self._extract_value(image_text, "Product picture:")

        detail_cell = self._get_cell(tables[2], 0, 1)
        category = self._extract_value(detail_cell, "Category:")
        subcategory = self._extract_value(detail_cell, "Subcategory:")
        critical_item = self._extract_value(detail_cell, "Critical item:")
        dangerous_goods = self._extract_value(detail_cell, "Dangerous goods:")
        description = self._extract_value(detail_cell, "Short description:")

        if critical_item:
            metadata["critical_item"] = critical_item
        if dangerous_goods:
            metadata["dangerous_goods"] = dangerous_goods

        keywords: List[str] = []
        if category:
            keywords.append(category)
        if subcategory:
            keywords.append(subcategory)

        repo_cell = self._get_cell(tables[2], 1, 1)
        repo = self._extract_value(repo_cell, "Repository link:")

        dim_cell = self._get_cell(tables[2], 2, 0)
        dims_str = self._extract_value(dim_cell, "LxWxH:")
        use_type = self._extract_value(dim_cell, "Single/Multiple use:")
        solution_type = self._extract_value(dim_cell, "Permanent/Temporary solution:")
        if use_type:
            metadata["use_type"] = use_type
        if solution_type:
            metadata["solution_type"] = solution_type

        license_cell = self._get_cell(tables[2], 2, 1)
        license_str = self._extract_value(license_cell, "License:")

        readiness_cell = self._get_cell(tables[2], 3, 1)
        field_readiness = self._extract_value(readiness_cell, "Field readiness:")
        maker_readiness = self._extract_value(readiness_cell, "Maker readiness:")
        user_readiness = self._extract_value(readiness_cell, "User readiness:")
        tech_readiness = self._extract_value(readiness_cell, "Technology readiness:")
        risk_level = self._extract_value(readiness_cell, "Risk level:")
        if field_readiness:
            metadata["field_readiness"] = field_readiness
        if maker_readiness:
            metadata["maker_readiness"] = maker_readiness
        if user_readiness:
            metadata["user_readiness"] = user_readiness
        if risk_level:
            metadata["risk_level"] = risk_level

        justification_cell = self._get_cell(tables[2], 4, 0)
        intended_use = self._extract_value(
            justification_cell, "Justification of using 3D printed item:"
        )

        why_cell = self._get_cell(tables[2], 5, 0)
        why_3d = self._extract_value(why_cell, "Why should this item be 3D printed:")
        approval = self._extract_value(why_cell, "Approval required by:")
        if why_3d:
            metadata["justification_3d_print"] = why_3d
        if approval:
            metadata["approval_required_by"] = approval

        # --- Table 3: FIT ------------------------------------------------
        compat_cell = self._get_cell(tables[3], 0, 0)
        primary_compat = self._extract_value(compat_cell, "Primary compatibility:")
        accessories = self._extract_value(compat_cell, "Compatible accessories:")
        if primary_compat:
            metadata["primary_compatibility"] = primary_compat
        if accessories:
            metadata["compatible_accessories"] = accessories

        mfg_instr_cell = self._get_cell(tables[3], 1, 0)
        mfg_instr_text = self._extract_value(
            mfg_instr_cell, "Manufacturing instructions:"
        )

        detail_fit_cell = self._get_cell(tables[3], 2, 0)
        material_color = self._extract_value(detail_fit_cell, "Material and color:")
        other_materials = self._extract_value(
            detail_fit_cell, "List of other materials:"
        )
        printer_3d = self._extract_value(detail_fit_cell, "3D Printer:")
        slicer_settings = self._extract_value(detail_fit_cell, "Slicer settings:")
        post_proc = self._extract_value(
            detail_fit_cell, "Post processing instructions:"
        )
        assembly_instr = self._extract_value(detail_fit_cell, "Assembly instructions:")
        visual_insp = self._extract_value(detail_fit_cell, "Visual inspection:")
        dim_validation = self._extract_value(detail_fit_cell, "Dimensional validation:")
        tol_insp = self._extract_value(detail_fit_cell, "Tolerance inspection:")
        safety_valid = self._extract_value(detail_fit_cell, "Safety validation:")
        check_freq = self._extract_value(
            detail_fit_cell, "Regular product check frequency and responsibility:"
        )

        if other_materials:
            metadata["other_materials"] = other_materials
        if slicer_settings:
            metadata["slicer_settings"] = slicer_settings
        if post_proc:
            metadata["post_processing"] = post_proc
        if assembly_instr:
            metadata["assembly_instructions"] = assembly_instr
        if visual_insp:
            metadata["visual_inspection"] = visual_insp
        if dim_validation:
            metadata["dimensional_validation"] = dim_validation
        if tol_insp:
            metadata["tolerance_inspection"] = tol_insp
        if safety_valid:
            metadata["safety_validation"] = safety_valid
        if check_freq:
            metadata["check_frequency"] = check_freq

        # --- Table 4: FUNCTION -------------------------------------------
        function_cell = self._get_cell(tables[4], 0, 0)
        function_text = self._extract_value(
            function_cell,
            "Detailed description of the component/product/workflow and its use:",
        )

        cleaning_cell = self._get_cell(tables[4], 2, 0)
        cleaning = self._extract_value(
            cleaning_cell, "Cleaning and sanitization procedures:"
        )
        if cleaning:
            metadata["cleaning_procedures"] = cleaning

        packaging_cell = self._get_cell(tables[4], 4, 0)
        packaging = self._extract_value(
            packaging_cell, "Packaging and storing instructions:"
        )
        if packaging:
            metadata["packaging_instructions"] = packaging

        standards_cell = self._get_cell(tables[4], 6, 0)
        standards_text = self._extract_value(
            standards_cell, "Related links, standards, safety considerations:"
        )

        spaulding_cell = self._get_cell(tables[4], 7, 0)
        spaulding = self._extract_value(
            spaulding_cell, "Spaulding Classification (IPC):"
        )
        if spaulding:
            metadata["spaulding_classification"] = spaulding

        # --- Table 5: ATTACHMENTS ----------------------------------------
        designer_cell = self._get_cell(tables[5], 1, 0)
        designer = self._extract_value(designer_cell, "Designed by")

        design_date_cell = self._get_cell(tables[5], 1, 2)
        design_date_str = self._extract_value(design_date_cell, "Date:")

        approved_cell = self._get_cell(tables[5], 2, 0)
        approved_by = self._extract_value(approved_cell, "Product approved by")
        if approved_by:
            metadata["approved_by"] = approved_by

        approved_date_cell = self._get_cell(tables[5], 2, 2)
        approved_date = self._extract_value(approved_date_cell, "Date:")
        if approved_date:
            metadata["approved_date"] = approved_date

        tested_cell = self._get_cell(tables[5], 3, 0)
        tested_by = self._extract_value(tested_cell, "Product tested by")
        if tested_by:
            metadata["tested_by"] = tested_by

        tested_date_cell = self._get_cell(tables[5], 3, 2)
        tested_date = self._extract_value(tested_date_cell, "Date:")
        if tested_date:
            metadata["tested_date"] = tested_date

        # ------------------------------------------------------------------
        # Build the OKHManifest
        # ------------------------------------------------------------------

        # License
        license_obj = License(hardware=license_str or None)

        # Licensor / designer
        licensor = designer or "Unknown"

        # Materials
        materials: List[MaterialSpec] = []
        if material_color:
            for mat_name in [m.strip() for m in material_color.split(",") if m.strip()]:
                materials.append(MaterialSpec(material_id="", name=mat_name))

        # Tool list (3D printers)
        tool_list: List[str] = []
        if printer_3d:
            tool_list = [t.strip() for t in printer_3d.split(",") if t.strip()]

        # Making instructions
        making_instructions: List[DocumentRef] = []
        if mfg_instr_text:
            making_instructions.append(
                DocumentRef(
                    title="Manufacturing Instructions",
                    path=mfg_instr_text,
                    type=DocumentationType.MAKING_INSTRUCTIONS,
                )
            )

        # Operating instructions (cleaning, packaging)
        operating_instructions: List[DocumentRef] = []
        if cleaning:
            operating_instructions.append(
                DocumentRef(
                    title="Cleaning and Sanitization",
                    path="cleaning-procedures",
                    type=DocumentationType.OPERATING_INSTRUCTIONS,
                    metadata={"content": cleaning},
                )
            )
        if packaging:
            operating_instructions.append(
                DocumentRef(
                    title="Packaging and Storing",
                    path="packaging-instructions",
                    type=DocumentationType.OPERATING_INSTRUCTIONS,
                    metadata={"content": packaging},
                )
            )

        # Standards
        standards_used: List[Standard] = []
        if standards_text:
            for std in [s.strip() for s in standards_text.split(";") if s.strip()]:
                standards_used.append(Standard(standard_title=std))

        # Version date
        version_date = None
        if design_date_str:
            try:
                version_date = date.fromisoformat(design_date_str.strip())
            except ValueError:
                logger.warning(f"Could not parse date: {design_date_str}")

        # Parse outer dimensions
        outer_dimensions = None
        if dims_str:
            outer_dimensions = self._parse_dimensions(dims_str)

        manifest = OKHManifest(
            title=title or "Untitled",
            version="1.0.0",
            license=license_obj,
            licensor=licensor,
            documentation_language="en",
            function=function_text or description or "",
            repo=repo or None,
            description=description or None,
            intended_use=intended_use or None,
            keywords=keywords,
            image=image or None,
            development_stage=dev_stage or None,
            technology_readiness_level=tech_readiness or None,
            health_safety_notice=standards_text or None,
            materials=materials,
            tool_list=tool_list,
            making_instructions=making_instructions,
            operating_instructions=operating_instructions,
            standards_used=standards_used,
            version_date=version_date,
            metadata=metadata,
        )

        # Set manufacturing specs if dimensions were parsed
        if outer_dimensions:
            from ..models.okh import ManufacturingSpec

            manifest.manufacturing_specs = ManufacturingSpec(
                outer_dimensions=outer_dimensions
            )

        logger.info(f"OKH manifest parsed from datasheet: {manifest.title}")
        return manifest

    # ------------------------------------------------------------------
    # Helper methods
    # ------------------------------------------------------------------

    @staticmethod
    def _set_cell(table, row: int, col: int, text: str) -> None:
        """Safely set the text content of a table cell."""
        try:
            cell = table.rows[row].cells[col]
            # Clear existing content
            for paragraph in cell.paragraphs:
                paragraph.text = ""
            if cell.paragraphs:
                cell.paragraphs[0].text = text
            else:
                cell.add_paragraph(text)
        except (IndexError, AttributeError) as exc:
            logger.warning(f"Could not set cell [{row},{col}]: {exc}")

    @staticmethod
    def _get_cell(table, row: int, col: int) -> str:
        """Safely get the text content of a table cell."""
        try:
            return table.rows[row].cells[col].text.strip()
        except (IndexError, AttributeError):
            return ""

    @staticmethod
    def _extract_value(cell_text: str, label: str) -> str:
        """Extract a value from cell text that follows a label.

        Handles patterns like:
          "Label: value"
          "Label:\\nvalue"

        The value is considered to be everything after the label until the
        next known label (next line starting with a capitalised word
        followed by a colon) or end of string.
        """
        if not cell_text or label not in cell_text:
            return ""

        # Find the position after the label
        idx = cell_text.index(label) + len(label)
        rest = cell_text[idx:].strip()

        # Remove leading colon if present (labels sometimes end with ':')
        # but we may have already included it in the label
        if rest.startswith(":"):
            rest = rest[1:].strip()

        # Find the next label boundary (a line that looks like "SomeLabel:")
        # We split by newlines and take text until the next label-like line.
        lines = rest.split("\n")
        result_lines: List[str] = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            # The first line is always part of the value
            if i == 0:
                result_lines.append(stripped)
                continue
            # Check if this line looks like a new label
            if re.match(r"^[A-Z][a-zA-Z\s/]*:", stripped):
                break
            result_lines.append(stripped)

        return "\n".join(result_lines).strip()

    @staticmethod
    def _license_to_string(license_obj: License) -> str:
        """Convert a License object to a human-readable string."""
        parts = []
        if license_obj.hardware:
            parts.append(license_obj.hardware)
        if (
            license_obj.documentation
            and license_obj.documentation != license_obj.hardware
        ):
            parts.append(f"Docs: {license_obj.documentation}")
        if license_obj.software and license_obj.software != license_obj.hardware:
            parts.append(f"SW: {license_obj.software}")
        return ", ".join(parts) if parts else ""

    @staticmethod
    def _licensor_to_string(licensor) -> str:
        """Convert a licensor field to a string."""
        if isinstance(licensor, str):
            return licensor
        if isinstance(licensor, Person):
            return licensor.name
        if isinstance(licensor, list):
            names = []
            for item in licensor:
                if isinstance(item, str):
                    names.append(item)
                elif hasattr(item, "name"):
                    names.append(item.name)
            return ", ".join(names)
        if hasattr(licensor, "name"):
            return licensor.name
        return str(licensor) if licensor else ""

    @staticmethod
    def _get_doc_ref_text(doc_refs: List[DocumentRef], keyword: str) -> str:
        """Search a list of DocumentRef for one whose title contains keyword."""
        for ref in doc_refs:
            if keyword.lower() in ref.title.lower():
                return ref.metadata.get("content", ref.path)
        return ""

    @staticmethod
    def _parse_dimensions(dims_str: str) -> Optional[Dict[str, Any]]:
        """Parse a dimension string like '100x50x25' into a dict."""
        dims_str = dims_str.strip()
        # Try to match patterns like "100x50x25" or "100 x 50 x 25"
        match = re.match(
            r"(\d+(?:\.\d+)?)\s*[xX×]\s*(\d+(?:\.\d+)?)\s*[xX×]\s*(\d+(?:\.\d+)?)",
            dims_str,
        )
        if match:
            return {
                "length": float(match.group(1)),
                "width": float(match.group(2)),
                "height": float(match.group(3)),
            }
        return None
